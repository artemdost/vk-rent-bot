"""
Сервис для извлечения текста из различных типов документов.
Поддерживает PDF, DOCX, TXT и изображения.
"""
import os
import tempfile
import requests
from typing import Optional, List, Any
import logging
import base64
import io

# Импорты для работы с документами
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 не установлен. PDF файлы не будут обрабатываться.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx не установлен. DOCX файлы не будут обрабатываться.")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    # OCR библиотеки опциональны - изображения можно обрабатывать через Deepseek API


class DocumentParser:
    """Парсер документов различных форматов."""

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    SUPPORTED_IMAGE_FORMATS = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']
    SUPPORTED_DOCUMENT_FORMATS = ['.pdf', '.docx', '.txt']

    def __init__(self):
        """Инициализация парсера документов."""
        self.logger = logging.getLogger(__name__)

    async def download_file(self, url: str) -> Optional[bytes]:
        """
        Скачивает файл по URL.

        Args:
            url: URL файла

        Returns:
            Содержимое файла в байтах или None при ошибке
        """
        try:
            response = requests.get(url, timeout=30, stream=True)
            response.raise_for_status()

            # Проверка размера файла
            content_length = response.headers.get('Content-Length')
            if content_length and int(content_length) > self.MAX_FILE_SIZE:
                self.logger.error(f"Файл слишком большой: {content_length} байт")
                return None

            # Загрузка с ограничением размера
            content = b""
            for chunk in response.iter_content(chunk_size=8192):
                content += chunk
                if len(content) > self.MAX_FILE_SIZE:
                    self.logger.error("Файл превышает максимальный размер")
                    return None

            return content

        except requests.RequestException as e:
            self.logger.error(f"Ошибка загрузки файла: {e}")
            return None

    def extract_text_from_pdf(self, content: bytes) -> Optional[str]:
        """
        Извлекает текст из PDF файла.

        Args:
            content: Содержимое PDF файла

        Returns:
            Извлеченный текст или None при ошибке
        """
        if not PDF_AVAILABLE:
            return None

        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"

            return text.strip()

        except Exception as e:
            self.logger.error(f"Ошибка извлечения текста из PDF: {e}")
            return None

    def extract_text_from_docx(self, content: bytes) -> Optional[str]:
        """
        Извлекает текст из DOCX файла.

        Args:
            content: Содержимое DOCX файла

        Returns:
            Извлеченный текст или None при ошибке
        """
        if not DOCX_AVAILABLE:
            return None

        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            self.logger.error(f"Ошибка извлечения текста из DOCX: {e}")
            return None

    def extract_text_from_image(self, content: bytes) -> Optional[str]:
        """
        Извлекает текст из изображения с помощью OCR или возвращает base64 для Deepseek.

        Args:
            content: Содержимое изображения

        Returns:
            Извлеченный текст, base64 строка или None при ошибке
        """
        if not OCR_AVAILABLE:
            # Если OCR недоступен, конвертируем изображение в base64 для отправки в Deepseek
            try:
                import base64
                base64_image = base64.b64encode(content).decode('utf-8')
                return f"[IMAGE:{base64_image}]"
            except Exception as e:
                self.logger.error(f"Ошибка кодирования изображения в base64: {e}")
                return "[Изображение документа - не удалось обработать]"

        try:
            image = Image.open(io.BytesIO(content))

            # Конвертация в RGB если необходимо
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # OCR с русским языком
            text = pytesseract.image_to_string(image, lang='rus+eng')

            return text.strip()

        except Exception as e:
            self.logger.error(f"Ошибка OCR: {e}")
            return "[Не удалось распознать текст в изображении]"

    def extract_text_from_txt(self, content: bytes) -> Optional[str]:
        """
        Извлекает текст из TXT файла.

        Args:
            content: Содержимое TXT файла

        Returns:
            Извлеченный текст или None при ошибке
        """
        try:
            # Пробуем разные кодировки
            for encoding in ['utf-8', 'windows-1251', 'cp866']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # Если не удалось декодировать, пробуем с игнорированием ошибок
            return content.decode('utf-8', errors='ignore')

        except Exception as e:
            self.logger.error(f"Ошибка чтения TXT файла: {e}")
            return None

    async def extract_text_from_attachment(self, attachment: Any) -> Optional[str]:
        """
        Извлекает текст из вложения VK.

        Args:
            attachment: Объект вложения из VK API (VKBottle Pydantic model)

        Returns:
            Извлеченный текст или None при ошибке
        """
        try:
            # VKBottle использует Pydantic модели, получаем тип через атрибут
            attachment_type = getattr(attachment, 'type', None)

            # Обработка документов
            if attachment_type == 'doc':
                doc = getattr(attachment, 'doc', None)
                if doc:
                    url = getattr(doc, 'url', None)
                    ext = getattr(doc, 'ext', '').lower()

                if not url:
                    return None

                # Скачиваем файл
                content = await self.download_file(url)
                if not content:
                    return None

                # Извлекаем текст в зависимости от типа
                if ext == 'pdf' or url.lower().endswith('.pdf'):
                    return self.extract_text_from_pdf(content)
                elif ext == 'docx' or url.lower().endswith('.docx'):
                    return self.extract_text_from_docx(content)
                elif ext == 'txt' or url.lower().endswith('.txt'):
                    return self.extract_text_from_txt(content)
                else:
                    self.logger.warning(f"Неподдерживаемый тип документа: {ext}")
                    return None

            # Обработка изображений
            elif attachment_type == 'photo':
                photo = getattr(attachment, 'photo', None)
                if not photo:
                    return None

                # Ищем изображение максимального размера
                max_size_url = None

                # VKBottle photo object has sizes attribute
                sizes = getattr(photo, 'sizes', None)
                if sizes:
                    # Сортируем по размеру (ширина * высота)
                    sorted_sizes = sorted(sizes, key=lambda x: getattr(x, 'width', 0) * getattr(x, 'height', 0), reverse=True)
                    if sorted_sizes:
                        max_size_url = getattr(sorted_sizes[0], 'url', None)

                if not max_size_url:
                    return None

                # Скачиваем изображение
                content = await self.download_file(max_size_url)
                if not content:
                    return None

                # Извлекаем текст с помощью OCR
                return self.extract_text_from_image(content)

            else:
                self.logger.warning(f"Неподдерживаемый тип вложения: {attachment_type}")
                return None

        except Exception as e:
            self.logger.error(f"Ошибка обработки вложения: {e}")
            return None

    async def extract_all_texts(self, attachments: List[Any]) -> List[str]:
        """
        Извлекает текст из всех вложений.

        Args:
            attachments: Список вложений из VK API (VKBottle Pydantic models)

        Returns:
            Список извлеченных текстов
        """
        texts = []

        for attachment in attachments:
            text = await self.extract_text_from_attachment(attachment)
            if text:
                texts.append(text)

        return texts

    def get_file_type_description(self, attachment: Any) -> str:
        """
        Получает описание типа файла для пользователя.

        Args:
            attachment: Объект вложения из VK API (VKBottle Pydantic model)

        Returns:
            Описание типа файла
        """
        attachment_type = getattr(attachment, 'type', None)

        if attachment_type == 'doc':
            doc = getattr(attachment, 'doc', None)
            if doc:
                ext = getattr(doc, 'ext', '').upper()
                title = getattr(doc, 'title', 'документ')
                return f"{ext} документ '{title}'"
            return "документ"
        elif attachment_type == 'photo':
            return "изображение"
        else:
            return "файл"